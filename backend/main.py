import json
import os
import traceback
import uuid
from datetime import datetime, timedelta
from io import BytesIO
from typing import Any, Dict, List, Optional, Union

import anthropic
import aspose.words as aw
import env_variables
import uvicorn
from cache_manager import (add_to_filename_map, get_cached_data,
                           get_cached_file_id, save_cached_data)
from docx import Document
from docx.shared import Inches
from dotenv import load_dotenv
from fastapi import (Depends, FastAPI, Header, HTTPException, Query, Response,
                     UploadFile, status)
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from generate_raw_data import (get_areas_to_target_data, get_raw_data,
                               get_strengths_data)
from generate_report_llm import (extract_employee_info, process_prompts,
                                 read_file_content,
                                 transform_content_to_report_format)
from jose import JWTError, jwt
from passlib.context import CryptContext
from process_pdf import AssessmentProcessor
from prompt_loader import (format_area_content_prompt,
                           format_next_steps_prompt,
                           format_reflection_points_prompt,
                           format_strength_content_prompt, load_prompt)
from pydantic import BaseModel
from report_generation import (create_360_feedback_report,
                               create_360_feedback_report_for_word)
from state import files_store
from utils.jwt_utils import verify_clerk_token
from utils.postgreSql_uitls import get_db_connection
from utils.validate_envs import validate_required_env

from routers.advice import router as advice_routers
from routers.feedback import router as feedback_routers
from routers.file import router as file_routers
from routers.snapshot import router as snapshot_routers
from auth.user import User, get_current_user
from sqlalchemy.orm import Session
from db.core import get_db
from db.feedback import get_cached_feedback



import json
import os

def read_filename_map():
    try:
        # Define the path to the filename_map.json file
        CACHE_DIR = "../data/cache"
        filepath = os.path.join(CACHE_DIR, "filename_map.json")
        
        # Open and read the file
        with open(filepath, 'r') as file:
            filename_map = json.load(file)
            
        print(f"Successfully loaded filename map with {len(filename_map)} entries")
        return filename_map
    
    except FileNotFoundError:
        print(f"Error: File not found at {filepath}")
        return {}
    except json.JSONDecodeError:
        print(f"Error: Invalid JSON format in {filepath}")
        return {}
    except Exception as e:
        print(f"Error reading filename map: {str(e)}")
        return {}

# Example usage:


# Initialize license object
license = aw.License()

# Set license
try:
    # Method 1: Load license from file
    license.set_license("Aspose.WordsforPythonvia.NET.lic")
    print("License set successfully!")
except Exception as e:
    print(f"Error setting license: {str(e)}")

api_key = env_variables.ANTHROPIC_API_KEY
assessment_processor = AssessmentProcessor(api_key)


# Password hashing
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

# User model
# class User(BaseModel):
#     username: str
#     hashed_password: str


# Simple authentication response
class AuthResponse(BaseModel):
    success: bool
    username: str


# User database (in-memory for simplicity)
# In a production environment, this would be a database
users_db = {"Tom": {"username": "Tom", "hashed_password": pwd_context.hash("Tom@1234")}}

# Track authenticated users (in-memory for simplicity)
authenticated_users = set()

# Authentication functions
def verify_password(plain_password, hashed_password):
    return pwd_context.verify(plain_password, hashed_password)


def get_user(username: str):
    if username in users_db:
        user_dict = users_db[username]
        return User(**user_dict)
    return None


def authenticate_user(username: str, password: str):
    user = get_user(username)
    if not user:
        return False
    if not verify_password(password, user.hashed_password):
        return False
    return user


# Simple authentication check
# async def get_current_user(authorization: Optional[str] = Header(None)):
#     if not authorization or not authorization.startswith("Bearer "):
#         raise HTTPException(
#             status_code=401, detail="Missing or invalid Authorization header"
#         )

#     token = authorization.split("Bearer ")[1]
#     user_claims = verify_clerk_token(token)
#     user_id = user_claims.get("sub")
#     if not user_id:
#         user_id = (
#             user_claims.get("id")
#             or user_claims.get("user_id")
#             or user_claims.get("userId")
#         )
#     if user_id:
#         user_claims["user_id"] = user_id
#         return user_claims
#     return user_claims


app = FastAPI()

# Configure CORS

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000",
        "https://ff2c-34-202-149-23.ngrok-free.app",
        "http://34.202.149.23:3000",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Authentication endpoints
@app.post("/api/login", response_model=AuthResponse)
async def login_for_access_token(form_data: OAuth2PasswordRequestForm = Depends()):
    user = authenticate_user(form_data.username, form_data.password)
    if not user:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect username or password",
        )

    # Add user to authenticated users
    authenticated_users.add(user.username)

    return {"success": True, "username": user.username}


@app.get("/api/users/me")
async def read_users_me(username: str = Header(None)):
    if not username or username not in authenticated_users:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Not authenticated",
        )
    return {"username": username}


# Create uploads directory if it doesn't exist
UPLOAD_DIR = "../data/uploads"
SAVE_DIR = "../data/processed_assessments"
OUTPUT_DIR = "../data/output"
REPORT_DIR = "../data/generated_reports"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(SAVE_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(REPORT_DIR, exist_ok=True)

# Data models
class NextStepPoint(BaseModel):
    main: str
    sub_points: List[str]


class InterviewAnalysis(BaseModel):
    name: str
    date: str
    strengths: Dict[str, str]
    areas_to_target: Dict[str, str]
    next_steps: List[Union[str, NextStepPoint]]


class GenerateContentRequest(BaseModel):
    heading: str
    file_id: str
    existing_content: Optional[str] = None


class GenerateNextStepsRequest(BaseModel):
    areas_to_target: Dict[str, str]
    file_id: str
    executive_transcript: Optional[str] = None


class SortEvidenceRequest(BaseModel):
    file_id: str
    headings: List[str]


def get_name_from_report(file_id: str) -> str:
    report_file_path = os.path.join(REPORT_DIR, f"{file_id}_report.json")
    if not os.path.exists(report_file_path):
        raise HTTPException(
            status_code=404,
            detail="Report not found. Please generate the report first.",
        )

    with open(report_file_path, "r") as f:
        report_data = json.load(f)
    return report_data.get("name", "")


@app.post("/api/old/upload_file", tags=["old"])
async def upload_file(
    file: UploadFile,
    use_cache: bool = Query(
        True, description="Whether to use cached results if available"
    ),
):
    try:
        # Check if we have this file cached
        if use_cache:
            cached_file_id = get_cached_file_id(file.filename, use_cache=True)
            if cached_file_id:
                print(f"Using cached file ID {cached_file_id} for {file.filename}")
                # Check if the file still exists
                file_path = os.path.join(UPLOAD_DIR, f"{cached_file_id}.pdf")
                if os.path.exists(file_path):
                    # Update files_store if needed
                    if cached_file_id not in files_store:
                        files_store[cached_file_id] = {
                            "file_path": file_path,
                            "original_name": file.filename,
                        }
                    return {"file_id": cached_file_id}

        # If not cached or cache disabled, process normally
        file_id = str(uuid.uuid4())
        file_path = os.path.join(UPLOAD_DIR, f"{file_id}.pdf")

        # Save file to disk
        with open(file_path, "wb") as f:
            content = await file.read()
            f.write(content)

        # Process the file
        (
            stakeholder_feedback,
            executive_interview,
        ) = assessment_processor.process_assessment_with_executive(file_path, SAVE_DIR)

        # Store file info
        files_store[file_id] = {"file_path": file_path, "original_name": file.filename}

        # Add to cache mapping
        # breakpoint()
        add_to_filename_map(file.filename, file_id)

        return {"file_id": file_id}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/generate_report/{file_id}")
async def generate_report(
    file_id: str,
    use_cache: bool = Query(
        True, description="Whether to use cached results if available"
    ),
):  
    files_store = read_filename_map()
    if file_id not in files_store:
        raise HTTPException(status_code=404, detail="File not found")

    try:
        # Check cache first
        if use_cache:
            cached_data = get_cached_data("reports", file_id)
            if cached_data:
                print(f"Using cached report for file ID {file_id}")
                return cached_data

        # If not cached or cache disabled, process normally
        self_transcript = SAVE_DIR + "/executive_" + file_id + ".txt"
        others_transcript = SAVE_DIR + "/filtered_" + file_id + ".txt"

        feedback_content = read_file_content(others_transcript)
        executive_interview = read_file_content(self_transcript)

        system_prompt = ""

        results = await process_prompts(
            feedback_content, executive_interview, api_key, system_prompt
        )
        name_data = extract_employee_info(UPLOAD_DIR + "/" + file_id + ".pdf", api_key)

        employee_name = name_data.get("employee_name", "")
        report_date = name_data.get("report_date", "")

        formatted_data = transform_content_to_report_format(
            results, employee_name, report_date
        )

        # Save the formatted data to a file
        report_file_path = os.path.join(REPORT_DIR, f"{file_id}_report.json")
        with open(report_file_path, "w") as f:
            json.dump(formatted_data, f)

        # Save to cache
        save_cached_data("reports", file_id, formatted_data)

        return formatted_data
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

#Not being used anywhere
@app.get("/api/get_raw_data/{file_id}")
async def get_raw_data_endpoint(
    file_id: str,
    use_cache: bool = Query(
        True, description="Whether to use cached results if available"
    ),
):
    try:
        # Check cache first
        if use_cache:
            cached_data = get_cached_data("raw_data", file_id)
            if cached_data:
                print(f"Using cached raw data for file ID {file_id}")
                return cached_data

        # If not cached or cache disabled, process normally
        # Load the generated report
        report_file_path = os.path.join(REPORT_DIR, f"{file_id}_report.json")
        if not os.path.exists(report_file_path):
            raise HTTPException(
                status_code=404,
                detail="Report not found. Please generate the report first.",
            )

        with open(report_file_path, "r") as f:
            report_data = json.load(f)

        # Get the transcript
        transcript_path = os.path.join(SAVE_DIR, f"filtered_{file_id}.txt")
        if not os.path.exists(transcript_path):
            raise HTTPException(status_code=404, detail="Transcript not found.")

        with open(transcript_path, "r") as f:
            transcript = f.read()

        # Extract strengths and areas_to_target from the report data
        strengths = report_data.get("strengths", {})
        areas_to_target = report_data.get("areas_to_target", {})

        # Get raw data
        # For strengths analysis
        strengths_data = get_strengths_data(transcript, strengths, api_key)

        # For areas to target analysis
        areas_data = get_areas_to_target_data(transcript, areas_to_target, api_key)

        raw_data = {}
        raw_data.update(strengths_data)
        raw_data.update(areas_data)

        # Save to cache
        save_cached_data("raw_data", file_id, raw_data)

        return raw_data

    except Exception as e:
        print(f"Error in get_raw_data_endpoint: {str(e)}")
        import traceback

        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))


# @app.post("/api/dump_word")
# async def generate_word_document(analysis: InterviewAnalysis):
#     try:
#         # Create a new Word document
#         doc = Document()

#         # Add title
#         doc.add_heading('Interview Analysis Report', 0)

#         # Add basic info
#         doc.add_heading('Basic Information', level=1)
#         doc.add_paragraph(f'Name: {analysis.name}')
#         doc.add_paragraph(f'Date: {analysis.date}')

#         # Add strengths
#         doc.add_heading('Strengths', level=1)
#         for title, content in analysis.strengths.items():
#             doc.add_heading(title, level=2)
#             doc.add_paragraph(content)

#         # Add areas to target
#         doc.add_heading('Areas to Target', level=1)
#         for title, content in analysis.areas_to_target.items():
#             doc.add_heading(title, level=2)
#             doc.add_paragraph(content)

#         # Add next steps
#         doc.add_heading('Next Steps', level=1)
#         for step in analysis.next_steps:
#             if isinstance(step, str):
#                 doc.add_paragraph(step, style='List Bullet')
#             else:
#                 doc.add_paragraph(step.main)
#                 for sub_point in step.sub_points:
#                     p = doc.add_paragraph(style='List Bullet 2')
#                     p.add_run(sub_point)

#         # Save the document
#         output_path = os.path.join(UPLOAD_DIR, "interview_analysis.docx")
#         doc.save(output_path)

#         # Return the document as a response
#         return FileResponse(
#             output_path,
#             media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#             filename="interview_analysis.docx"
#         )
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/dump_word")
async def generate_word_document(
    analysis: InterviewAnalysis, 
):
    try:
        # First generate PDF with Word-optimized bullet points
        output_path = os.path.join(OUTPUT_DIR, "temp.pdf")
        header_txt = analysis.name + " - Qualitative 360 Feedback"
        create_360_feedback_report_for_word(output_path, analysis, header_txt)

        # Convert PDF to DOCX
        docx_path = os.path.join(OUTPUT_DIR, "Output1.docx")
        doc = aw.Document(output_path)
        doc.save(docx_path)

        # Return the DOCX file
        return FileResponse(
            docx_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="interview_analysis.docx",
        )
    except Exception as e:
        print(f"Error in generate_word_document: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/dump_pdf")
async def generate_pdf_docuement(
    analysis: InterviewAnalysis, 
):
    output_path = OUTPUT_DIR + "/temp.pdf"
    header_txt = analysis.name + " - Qualitative 360 Feedback"
    create_360_feedback_report(output_path, analysis, header_txt)
    return FileResponse(
        output_path, media_type="application/pdf", filename="interview_analysis.pdf"
    )


@app.post("/api/upload_updated_report")
async def upload_updated_report(
    file: UploadFile, 
):
    try:
        # Validate file type
        if not file.filename.endswith(".docx"):
            raise HTTPException(
                status_code=400, detail="Only Word documents (.docx) are allowed"
            )

        # Read docx content
        content = await file.read()
        doc = Document(BytesIO(content))
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

        # Use Claude to parse the content
        client = anthropic.Anthropic(api_key=api_key)
        # Load and format prompt
        prompt = load_prompt("upload_updated_report.txt").format(text=text)

        response = client.messages.create(
            model="claude-3-7-sonnet-latest",
            max_tokens=4000,
            temperature=0,
            messages=[{"role": "user", "content": prompt}],
        )

        # Parse response with better error handling
        response_text = response.content[0].text
        try:
            # Try direct JSON parsing first
            parsed_data = json.loads(response_text)
        except json.JSONDecodeError:
            # Try to extract JSON if surrounded by other text
            try:
                start = response_text.find("{")
                end = response_text.rfind("}") + 1
                if start >= 0 and end > 0:
                    json_str = response_text[start:end]
                    parsed_data = json.loads(json_str)
                else:
                    raise ValueError("No JSON content found in response")
            except Exception as e:
                print(f"Error parsing response: {str(e)}")
                print(f"Raw response: {response_text}")
                raise HTTPException(
                    status_code=500,
                    detail="Failed to parse AI response into valid JSON",
                )

        return parsed_data

    except Exception as e:
        print(f"Error in upload_updated_report: {str(e)}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/cleanup/{file_id}")
async def cleanup_file(file_id: str):
    if file_id in files_store:
        try:
            os.remove(files_store[file_id]["file_path"])
            del files_store[file_id]
            return {"message": "File cleaned up successfully"}
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))
    raise HTTPException(status_code=404, detail="File not found")


async def generate_reflection_points(
    feedback_transcript: str, executive_transcript: str
) -> dict:
    """Generate reflection points and context summary using both transcripts."""
    try:
        client = anthropic.Anthropic(api_key=api_key)
        prompt = format_reflection_points_prompt(
            feedback_transcript, executive_transcript
        )

        response = client.messages.create(
            model="claude-3-7-sonnet-latest",
            max_tokens=1000,
            temperature=0,
            messages=[{"role": "user", "content": prompt}],
        )

        # Extract JSON from response
        # print(response.content[0].text)
        response_text = response.content[0].text.replace("\n", "").replace("\\", "")
        try:
            result = json.loads(response_text)
        except json.JSONDecodeError:
            try:
                start = response_text.find("{")
                end = response_text.rfind("}") + 1
                if start >= 0 and end > 0:
                    json_str = response_text[start:end]
                    result = json.loads(json_str)
                else:
                    raise ValueError("No JSON content found in response")
            except Exception as e:
                print(f"Error parsing reflection points response: {str(e)}")
                print(f"Raw response: {response_text}")
                raise HTTPException(
                    status_code=500, detail="Failed to parse reflection points response"
                )

        return result
    except Exception as e:
        print(f"Error generating reflection points: {str(e)}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))


@app.options("/api/generate_next_steps")
async def generate_next_steps_options():
    headers = {
        "Access-Control-Allow-Origin": "http://localhost:3000",
        "Access-Control-Allow-Methods": "POST, OPTIONS",
        "Access-Control-Allow-Headers": "*",
        "Access-Control-Allow-Credentials": "true",
    }
    return Response(content="", headers=headers)


@app.post("/api/generate_next_steps")
async def generate_next_steps(
    request: dict,
):
    try:
        print("\n=== Generate Next Steps Request ===")
        file_id = request.get("file_id")
        areas_to_target = request.get("areas_to_target")
        if not file_id:
            raise HTTPException(status_code=400, detail="file_id is required")
        if not areas_to_target:
            raise HTTPException(status_code=400, detail="areas_to_target is required")

        print(f"File ID: {file_id}")
        print(f"Areas to Target: {json.dumps(areas_to_target, indent=2)}")
        print(f"Number of areas: {len(areas_to_target)}")

        # Load both transcripts using file ID
        feedback_path = f"../data/processed_assessments/filtered_{file_id}.txt"
        executive_path = f"../data/processed_assessments/executive_{file_id}.txt"

        if not os.path.exists(feedback_path):
            raise HTTPException(
                status_code=404,
                detail=f"Feedback transcript not found at {feedback_path}",
            )

        with open(feedback_path, "r") as f:
            feedback_transcript = f.read()

        executive_transcript = ""
        if os.path.exists(executive_path):
            with open(executive_path, "r") as f:
                executive_transcript = f.read()

        # Get reflection points first
        reflection_points = await generate_reflection_points(
            feedback_transcript, executive_transcript
        )
        print("\n=== Reflection Points ===")
        print(json.dumps(reflection_points, indent=2))

        # Format areas to target for action steps prompt
        areas_text = "\n".join(
            [f"{title}: {content}" for title, content in areas_to_target.items()]
        )
        print("\n=== Formatted Areas Text ===")
        print(areas_text)

        # Generate content using Claude
        client = anthropic.Anthropic(api_key=api_key)
        name = get_name_from_report(file_id)
        prompt = format_next_steps_prompt(name, areas_text, feedback_transcript)

        response = client.messages.create(
            model="claude-3-7-sonnet-latest",
            max_tokens=2000,
            temperature=0,
            messages=[{"role": "user", "content": prompt}],
        )

        # Extract JSON from response
        response_text = response.content[0].text.replace("\n", "").replace("\\", "")
        print("\n=== Claude Response ===")
        print(response_text)
        try:
            # Try to parse the entire response as JSON first
            result = json.loads(response_text)
        except json.JSONDecodeError:
            # If that fails, try to extract JSON from the text
            try:
                # Find JSON-like content between curly braces
                start = response_text.find("{")
                end = response_text.rfind("}") + 1
                if start >= 0 and end > 0:
                    json_str = response_text[start:end]
                    result = json.loads(json_str)
                else:
                    raise ValueError("No JSON content found in response")
            except Exception as e:
                print(f"Error parsing response: {str(e)}")
                print(f"Raw response: {response_text}")
                raise HTTPException(
                    status_code=500,
                    detail="Failed to parse AI response into valid JSON",
                )

        # Ensure the response has the expected structure
        if "next_steps" not in result:
            result = {"next_steps": result}

        # Combine reflection points and action steps
        if "reflection_prompts" in reflection_points:
            prompts = reflection_points["reflection_prompts"]
            result["next_steps"] = [
                {
                    "main": prompts["discussion_prompt"],
                    "sub_points": prompts["bullet_points"],
                },
                prompts["context_summary"],
                *result["next_steps"],
            ]

        print("\n=== Final Result ===")
        print(json.dumps(result, indent=2))
        return result
    except Exception as e:
        print(f"Error in generate_next_steps: {str(e)}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/generate_area_content")
async def generate_area_content(
    request: GenerateContentRequest 
):
    try:
        print(
            f"[Area Content] Received request - heading: '{request.heading}', file_id: {request.file_id}, has_existing_content: {request.existing_content}"
        )
        # print("Existing content: ",request.existing_content)
        # Load transcript using file ID
        feedback_path = f"../data/processed_assessments/filtered_{request.file_id}.txt"
        if not os.path.exists(feedback_path):
            raise HTTPException(
                status_code=404,
                detail=f"Feedback transcript not found at {feedback_path}",
            )

        with open(feedback_path, "r") as f:
            feedback_transcript = f.read()

        # Generate content using Claude
        client = anthropic.Anthropic(api_key=api_key)
        name = get_name_from_report(request.file_id)
        prompt = format_area_content_prompt(
            name, request.heading, feedback_transcript, request.existing_content
        )
        print(prompt)
        response = client.messages.create(
            model="claude-3-7-sonnet-latest",
            max_tokens=1000,
            temperature=0,
            messages=[{"role": "user", "content": prompt}],
        )

        return {"content": response.content[0].text}
    except Exception as e:
        print(f"Error in generate_area_content: {str(e)}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))


def repair_json(json_str):
    """
    Advanced JSON repair algorithm that can handle nested structures.
    
    Args:
        json_str: The JSON string to repair
        
    Returns:
        Repaired JSON string
    """
    # Track nesting levels and string context
    in_string = False
    escape_next = False
    brace_stack = []
    bracket_stack = []
    
    # Process character by character
    repaired = []
    for i, char in enumerate(json_str):
        repaired.append(char)
        
        # Handle string context
        if char == '"' and not escape_next:
            in_string = not in_string
        
        # Track escape sequences
        if char == '\\' and not escape_next:
            escape_next = True
        else:
            escape_next = False
            
        # Only track structure outside of strings
        if not in_string:
            # Track array nesting
            if char == '[':
                bracket_stack.append(i)
            elif char == ']' and bracket_stack:
                bracket_stack.pop()
                
            # Track object nesting
            if char == '{':
                brace_stack.append(i)
            elif char == '}' and brace_stack:
                brace_stack.pop()
    
    # Fix unterminated strings
    if in_string:
        repaired.append('"')
    
    # Close any unclosed arrays
    while bracket_stack:
        bracket_stack.pop()
        repaired.append(']')
    
    # Close any unclosed objects
    while brace_stack:
        brace_stack.pop()
        repaired.append('}')
    
    return ''.join(repaired)


def parse_claude_response(response_text):
    """
    Enhanced JSON parsing with repair capability for Claude API responses.
    
    Args:
        response_text: The raw response text from Claude API
        
    Returns:
        Parsed JSON object
        
    Raises:
        Exception: If parsing fails after repair attempts
    """
    try:
        # First try direct parsing
        return json.loads(response_text)
    except json.JSONDecodeError as e:
        print(f"Initial JSON parsing failed: {str(e)}")
        try:
            # Find JSON array
            if response_text.strip().startswith('['):
                start = response_text.find("[")
                end = response_text.rfind("]") + 1
                if start >= 0 and end > 0:
                    json_str = response_text[start:end]
                    
                    # Basic repair for unterminated strings and other common issues
                    lines = json_str.split('\n')
                    repaired_lines = []
                    
                    for i, line in enumerate(lines):
                        # Check for unterminated strings in name field
                        if '"name": "' in line and not (line.strip().endswith('",') or line.strip().endswith('"')):
                            line = line.rstrip() + '",'
                        
                        # Check for unterminated strings in position field
                        if '"position": "' in line and not (line.strip().endswith('",') or line.strip().endswith('"')):
                            line = line.rstrip() + '",'
                        
                        # Check for unterminated strings in quote field
                        if '"quote": "' in line and not (line.strip().endswith('",') or line.strip().endswith('"')):
                            line = line.rstrip() + '",'
                        
                        # Handle unterminated isStrong field
                        if '"isStrong": ' in line and not any(line.strip().endswith(x) for x in ['true', 'false', 'true,', 'false,']):
                            if 'true' in line.lower():
                                line = line.rstrip() + 'true,'
                            else:
                                line = line.rstrip() + 'false,'
                        
                        # Check for missing commas between objects
                        if i < len(lines) - 1:
                            next_line = lines[i+1].strip()
                            if line.strip().endswith('}') and next_line.startswith('{'):
                                line = line.rstrip() + ','
                            elif line.strip().endswith(']') and next_line.startswith('['):
                                line = line.rstrip() + ','
                        
                        repaired_lines.append(line)
                    
                    repaired_json = '\n'.join(repaired_lines)
                    
                    try:
                        return json.loads(repaired_json)
                    except json.JSONDecodeError:
                        # If still failing, try a more aggressive approach
                        print("First repair attempt failed, trying more aggressive repair")
                        repaired_json = repaired_json.replace('}\n  {', '},\n  {')
                        repaired_json = repaired_json.replace(']\n  [', '],\n  [')
                        
                        try:
                            return json.loads(repaired_json)
                        except json.JSONDecodeError:
                            # If still failing, try the advanced repair algorithm
                            print("Second repair attempt failed, trying advanced repair")
                            repaired_json = repair_json(json_str)
                            return json.loads(repaired_json)
            
            # Find JSON object if not an array
            start = response_text.find("{")
            end = response_text.rfind("}") + 1
            if start >= 0 and end > 0:
                json_str = response_text[start:end]
                try:
                    return json.loads(json_str)
                except json.JSONDecodeError:
                    # Try advanced repair
                    repaired_json = repair_json(json_str)
                    return json.loads(repaired_json)
                
            raise ValueError("No JSON content found in response")
        except Exception as nested_e:
            print(f"Error in JSON repair: {str(nested_e)}")
            print(f"Original error: {str(e)}")
            print(f"Raw response: {response_text}")
            raise Exception(f"Failed to parse AI response into valid JSON: {str(e)}")


def merge_sorted_evidence(batch_results):
    """
    Merge multiple batches of sorted evidence into a single result.
    
    Args:
        batch_results: List of results from different batches
        
    Returns:
        Merged result with all evidence
    """
    if not batch_results:
        return []
        
    # Initialize with the structure from the first batch
    merged_result = []
    heading_map = {}  # To track headings we've already seen
    
    # Process each batch
    for batch in batch_results:
        for item in batch:
            heading = item["heading"]
            
            # If we've seen this heading before, append evidence
            if heading in heading_map:
                existing_item = merged_result[heading_map[heading]]
                
                # Add new evidence, avoiding duplicates
                existing_quotes = {e["quote"] for e in existing_item["evidence"]}
                for evidence in item["evidence"]:
                    if evidence["quote"] not in existing_quotes:
                        existing_item["evidence"].append(evidence)
                        existing_quotes.add(evidence["quote"])
            else:
                # New heading, add to result
                heading_map[heading] = len(merged_result)
                merged_result.append(item)
    
    return merged_result


@app.post("/api/sort-strengths-evidence")
async def sort_strengths_evidence(
    request: SortEvidenceRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    try:
        # Get previously generated strengths from the database
        user_id = current_user.user_id
        feedback_data = None
        
        if user_id:
            try:
                feedback_data = get_cached_feedback(user_id, request.file_id, db)
            except Exception as e:
                print(f"Error getting cached feedback: {str(e)}")
        
        if not feedback_data:
            # If no cached feedback data, use the transcript directly
            feedback_path = f"../data/processed_assessments/filtered_{request.file_id}.txt"
            if not os.path.exists(feedback_path):
                raise HTTPException(
                    status_code=404,
                    detail=f"Feedback transcript not found at {feedback_path}",
                )

            with open(feedback_path, "r") as f:
                transcript = f.read()
                
            # Initialize Claude client
            client = anthropic.Anthropic(api_key=api_key)
            
            # Load and format prompt
            sort_prompt = load_prompt("sort_evidence_strenght.txt")
            prompt = sort_prompt.format(
                strengths=transcript, headings="\n".join(request.headings)
            )
        else:
            # Use the previously generated strengths from the database
            strengths_data = feedback_data.get("strengths", {})
            
            # Process the strengths data to ensure it has the is_strong flag
            processed_strengths = {}
            for person, data in strengths_data.items():
                processed_feedback = []
                for feedback_item in data.get("feedback", []):
                    if isinstance(feedback_item, dict):
                        # New format with is_strong flag
                        processed_feedback.append({
                            "text": feedback_item.get("text", ""),
                            "is_strong": feedback_item.get("is_strong", False)
                        })
                    else:
                        # Handle legacy format (plain string)
                        processed_feedback.append({
                            "text": feedback_item,
                            "is_strong": False
                        })
                
                processed_strengths[person] = {
                    "role": data.get("role", ""),
                    "feedback": processed_feedback
                }
            
            # Process in batches
            all_results = []
            stakeholders = list(processed_strengths.keys())
            batch_size = 1  # Process 1 stakeholder at a time for maximum reliability
            
            # Count total feedback items for verification
            total_feedback_items = sum(len(data.get("feedback", [])) for data in processed_strengths.values())
            print(f"Total feedback items to process: {total_feedback_items}")
            
            # Initialize Claude client
            client = anthropic.Anthropic(api_key=api_key)
            
            # Load prompt template
            sort_prompt = load_prompt("sort_evidence_strenght.txt")
            
            # Process in batches
            for i in range(0, len(stakeholders), batch_size):
                batch_stakeholders = stakeholders[i:i+batch_size]
                batch_strengths = {k: processed_strengths[k] for k in batch_stakeholders if k in processed_strengths}
                
                # Count items in this batch
                batch_items = sum(len(data.get("feedback", [])) for data in batch_strengths.values())
                print(f"Processing batch {i//batch_size + 1} with {len(batch_stakeholders)} stakeholders and {batch_items} feedback items")
                
                # Convert batch data to JSON
                batch_json = json.dumps(batch_strengths, indent=2)
                
                # Format prompt for this batch
                prompt = sort_prompt.format(
                    strengths=batch_json, 
                    headings="\n".join(request.headings)
                )
                
                # Get sorted evidence from Claude for this batch
                response = client.messages.create(
                    model="claude-3-7-sonnet-latest",
                    max_tokens=2000,
                    temperature=0,
                    messages=[{"role": "user", "content": prompt}],
                )
                
                # Parse JSON response using our enhanced parser
                response_text = response.content[0].text
                
                try:
                    # Use our enhanced parser that can handle common JSON issues
                    batch_result = parse_claude_response(response_text)
                except Exception as e:
                    print(f"Error parsing response: {str(e)}")
                    print(f"Raw response: {response_text}")
                    raise HTTPException(
                        status_code=500,
                        detail="Failed to parse AI response into valid JSON",
                    )
                
                # Count evidence items in the result
                evidence_count = sum(len(item.get("evidence", [])) for item in batch_result)
                print(f"Batch {i//batch_size + 1} returned {evidence_count} evidence items")
                
                # Verify all items were included
                if evidence_count < batch_items:
                    print(f"WARNING: Batch {i//batch_size + 1} is missing items. Expected {batch_items}, got {evidence_count}")
                
                all_results.append(batch_result)
            
            # Merge results from all batches
            result = merge_sorted_evidence(all_results)
            
            # Verify total evidence count
            total_evidence = sum(len(item.get("evidence", [])) for item in result)
            print(f"Total evidence items after merging: {total_evidence}")
            
            if total_evidence < total_feedback_items:
                print(f"WARNING: Final result is missing items. Expected {total_feedback_items}, got {total_evidence}")
            
            return result

        # If we're here, we're processing the entire dataset at once (not in batches)
        # Get sorted evidence from Claude
        response = client.messages.create(
            model="claude-3-7-sonnet-latest",
            max_tokens=2000,
            temperature=0,
            messages=[{"role": "user", "content": prompt}],
        )

        # Parse JSON response
        response_text = response.content[0].text

        try:
            result = json.loads(response_text)
        except json.JSONDecodeError:
            try:
                start = response_text.find("[")
                end = response_text.rfind("]") + 1
                if start >= 0 and end > 0:
                    json_str = response_text[start:end]
                    result = json.loads(json_str)
                else:
                    raise ValueError("No JSON content found in response")
            except Exception as e:
                print(f"Error parsing response: {str(e)}")
                print(f"Raw response: {response_text}")
                raise HTTPException(
                    status_code=500,
                    detail="Failed to parse AI response into valid JSON",
                )

        return result
    except Exception as e:
        print(f"Error in sort_strengths_evidence: {str(e)}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))


def transform_headings(data):
    """
    Transform headings in a nested dictionary by:
    1. Making the first letter uppercase
    2. Making all other letters lowercase
    3. Adding a period at the end of each heading

    Args:
        data (dict): The input dictionary with nested headings

    Returns:
        dict: The transformed dictionary with modified headings
    """
    if not isinstance(data, dict):
        return data

    result = {}

    for key, value in data.items():
        if key == "evidence":
            # This is the evidence list, keep it as is
            result[key] = value
        elif isinstance(value, dict) and "evidence" in value:
            # This is a heading with evidence, transform the key
            transformed_key = key[0].upper() + key[1:].lower()
            transformed_key = transformed_key.replace(".", "")
            transformed_key = transformed_key + "."
            result[transformed_key] = value
        else:
            # This is another nested dictionary, recursively process it
            result[key] = transform_headings(value)

    return result


def transform_strength_evidence(sorted_result):
    """
    Transform sorted strength evidence from flat array to the expected frontend format.
    
    Args:
        sorted_result (list): The sorted evidence in flat array format
        
    Returns:
        dict: The transformed evidence in nested object format
    """
    transformed = {"leadershipQualities": {}}
    
    for item in sorted_result:
        heading = item["heading"]
        evidence = item["evidence"]
        
        # Map field names from new format to old format
        mapped_evidence = []
        for e in evidence:
            mapped_evidence.append({
                "feedback": e["quote"],
                "source": e["name"],
                "role": e["position"],
                "is_strong": e["isStrong"]
            })
        
        transformed["leadershipQualities"][heading] = {
            "evidence": mapped_evidence
        }
    
    return transformed


def transform_area_evidence(sorted_result):
    """
    Transform sorted area evidence from flat array to the expected frontend format.
    
    Args:
        sorted_result (list): The sorted evidence in flat array format
        
    Returns:
        dict: The transformed evidence in nested object format
    """
    transformed = {"developmentAreas": {}}
    
    # Default competency mappings based on heading
    competency_mappings = {
        "Strategic thinking.": ["Strategic thinking", "Business acumen", "Problem solving"],
        "Collaborative influence.": ["Influence", "Communication", "Positive relationships"],
        "Directive leadership.": ["Communication", "Influence", "Driving results"],
        "Team development.": ["Develops the team", "Inspiring the team", "Positive relationships"],
        "External presence.": ["Business acumen", "Influence", "Org savvy"],
        "Results orientation.": ["Driving results", "Execution", "Planning"],
        "Organizational presence.": ["Org savvy", "Influence", "Positive relationships"],
        "Talent acceleration.": ["Develops the team", "Learning agility", "Inspiring the team"],
        "Additional areas.": ["Other"]
    }
    
    for item in sorted_result:
        heading = item["heading"]
        evidence = item["evidence"]
        
        # Map field names from new format to old format
        mapped_evidence = []
        for e in evidence:
            mapped_evidence.append({
                "feedback": e["quote"],
                "source": e["name"],
                "role": e["position"],
                "is_strong": e["isStrong"]
            })
        
        # Get competency alignment from mapping or use defaults
        competencies = competency_mappings.get(heading, ["Communication", "Influence"])
        
        transformed["developmentAreas"][heading] = {
            "competencyAlignment": competencies,
            "evidence": mapped_evidence
        }
    
    return transformed


@app.get("/api/get_strength_evidences/{file_id}")
async def get_strength_evidences(
    file_id: str,
    numCompetencies: int,
    use_cache: bool = Query(
        True, description="Whether to use cached results if available"
    ),
    current_user: Dict[str, Any] = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    try:
        # Check cache first with parameters
        if use_cache:
            cached_data = get_cached_data(
                "strength_evidences", file_id, {"num_competencies": numCompetencies}
            )
            if cached_data:
                print(
                    f"Using cached strength evidences for file ID {file_id} with {numCompetencies} competencies"
                )
                return cached_data

        # Get previously generated strengths from the database
        user_id = current_user.user_id if current_user else None
        feedback_data = None
        
        if user_id:
            try:
                feedback_data = get_cached_feedback(user_id, file_id, db)
            except Exception as e:
                print(f"Error getting cached feedback: {str(e)}")
        
        if not feedback_data:
            # If not in database, we need to generate it first
            raise HTTPException(
                status_code=400,
                detail="Feedback data not found. Please generate feedback data first."
            )
            
        print(f"Step 1: Generating competency headings for strengths")
        
        # Step 1: Generate competency headings using the new simplified prompt
        strength_prompt = load_prompt("strength_headings.txt")
        prompt = strength_prompt.format(num_competencies=numCompetencies)
        
        # Generate headings using Claude
        client = anthropic.Anthropic(api_key=api_key)
        response = client.messages.create(
            model="claude-3-7-sonnet-latest",
            max_tokens=1000,
            temperature=0,
            messages=[{"role": "user", "content": prompt}],
        )
        
        # Parse JSON response
        response_text = response.content[0].text
        try:
            headings_result = parse_claude_response(response_text)
        except Exception as e:
            print(f"Error parsing headings response: {str(e)}")
            print(f"Raw response: {response_text}")
            raise HTTPException(
                status_code=500,
                detail="Failed to parse AI response into valid JSON",
            )
        
        # Extract headings as a simple list
        headings = headings_result.get("headings", [])
        print(f"Generated {len(headings)} headings: {headings}")
        
        # Step 2: Use sort_strengths_evidence to organize evidence under these headings
        print(f"Step 2: Sorting evidence under generated headings")
        sort_request = SortEvidenceRequest(file_id=file_id, headings=headings)
        sorted_result = await sort_strengths_evidence(sort_request, current_user, db)
        
        # Step 3: Transform the result to match the expected frontend format
        print(f"Step 3: Transforming result to match expected frontend format")
        transformed_result = transform_strength_evidence(sorted_result)
        
        # Save to cache with parameters
        save_cached_data(
            "strength_evidences", file_id, transformed_result, {"num_competencies": numCompetencies}
        )
        
        return transformed_result
    except Exception as e:
        print(f"Error in get_strength_evidences: {str(e)}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))




@app.get("/api/get_advice/{file_id}")
async def get_advice(
    file_id: str,
    use_cache: bool = Query(
        True, description="Whether to use cached results if available"
    ),
):
    try:
        # Check cache first
        if use_cache:
            cached_data = get_cached_data("advice", file_id)
            if cached_data:
                print(f"Using cached advice for file ID {file_id}")
                return cached_data

        # If not cached or cache disabled, process normally
        # Get the feedback transcript
        feedback_path = os.path.join(SAVE_DIR, f"filtered_{file_id}.txt")
        if not os.path.exists(feedback_path):
            raise HTTPException(status_code=404, detail="Feedback transcript not found")

        with open(feedback_path, "r") as f:
            feedback_transcript = f.read()

        # Load and format prompt
        advice_prompt = load_prompt("advice.txt")
        prompt = advice_prompt.format(feedback=feedback_transcript)

        # Generate analysis using Claude
        client = anthropic.Anthropic(api_key=api_key)
        response = client.messages.create(
            model="claude-3-7-sonnet-latest",
            max_tokens=3000,
            temperature=0,
            messages=[{"role": "user", "content": prompt}],
        )

        # Parse JSON response
        response_text = response.content[0].text
        try:
            result = json.loads(response_text)
        except json.JSONDecodeError:
            try:
                # Find JSON-like content between curly braces
                start = response_text.find("{")
                end = response_text.rfind("}") + 1
                if start >= 0 and end > 0:
                    json_str = response_text[start:end]
                    result = json.loads(json_str)
                else:
                    raise ValueError("No JSON content found in response")
            except Exception as e:
                print(f"Error parsing response: {str(e)}")
                print(f"Raw response: {response_text}")
                raise HTTPException(
                    status_code=500,
                    detail="Failed to parse AI response into valid JSON",
                )

        # Save to cache
        save_cached_data("advice", file_id, result)

        return result
    except Exception as e:
        print(f"Error in get_advice: {str(e)}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))


# Original get_feedback endpoint commented out
# @app.get("/api/get_feedback/{file_id}")

# Dummy endpoint for testing the frontend with the new feedback structure
@app.get("/api/old_get_feedback/{file_id}")
# async def get_feedback(file_id: str, use_cache: bool = Query(True)):
#     """
#     Dummy endpoint that returns test data with the updated feedback structure.
#     This includes the is_strong flag for each feedback item.
#     """
#     # Create dummy data with the new structure
#     dummy_data = {
#         "strengths": {
#             "Manager": {
#                 "role": "Direct Manager",
#                 "feedback": [
#                     {"text": "Excellent communication skills with the team.", "is_strong": True},
#                     {"text": "Good at delegating tasks appropriately.", "is_strong": False},
#                     {"text": "Outstanding leadership during the project crisis.", "is_strong": True},
#                     {"text": "Handles pressure well.", "is_strong": False}
#                 ]
#             },
#             "Peer": {
#                 "role": "Team Member",
#                 "feedback": [
#                     {"text": "Always willing to help others.", "is_strong": False},
#                     {"text": "Exceptional problem-solving abilities.", "is_strong": True},
#                     {"text": "Good team player.", "is_strong": False}
#                 ]
#             }
#         },
#         "areas_to_target": {
#             "Manager": {
#                 "role": "Direct Manager",
#                 "feedback": [
#                     {"text": "Could improve time management skills.", "is_strong": False},
#                     {"text": "Significant issues with documentation quality.", "is_strong": True},
#                     {"text": "Sometimes misses deadlines.", "is_strong": False}
#                 ]
#             },
#             "Peer": {
#                 "role": "Team Member",
#                 "feedback": [
#                     {"text": "Needs to speak up more in meetings.", "is_strong": False},
#                     {"text": "Major communication gaps with other departments.", "is_strong": True},
#                     {"text": "Could be more proactive in identifying issues.", "is_strong": False}
#                 ]
#             }
#         }
#     }
    
#     return dummy_data
async def get_feedback(
    file_id: str,
    use_cache: bool = Query(
        True, description="Whether to use cached results if available"
    ),
):
    try:
        # Check cache first
        if use_cache:
            cached_data = get_cached_data("feedback", file_id)
            if cached_data:
                print(f"Using cached feedback for file ID {file_id}")
                return cached_data

        # If not cached or cache disabled, process normally
        # Get the feedback transcript
        feedback_path = os.path.join(SAVE_DIR, f"filtered_{file_id}.txt")
        if not os.path.exists(feedback_path):
            print(feedback_path)
            raise HTTPException(status_code=404, detail="Feedback transcript not found")

        with open(feedback_path, "r") as f:
            feedback_transcript = f.read()

        # Load and format prompts
        strengths_prompt = load_prompt("feedback_strengths.txt")
        areas_prompt = load_prompt("feedback_areas.txt")

        # Initialize Claude client
        client = anthropic.Anthropic(api_key=api_key)

        # Get strengths analysis
        strengths_response = client.messages.create(
            model="claude-3-7-sonnet-latest",
            max_tokens=3000,
            temperature=0,
            messages=[
                {
                    "role": "user",
                    "content": strengths_prompt.format(feedback=feedback_transcript),
                }
            ],
        )

        # Get areas analysis
        areas_response = client.messages.create(
            model="claude-3-7-sonnet-latest",
            max_tokens=3000,
            temperature=0,
            messages=[
                {
                    "role": "user",
                    "content": areas_prompt.format(feedback=feedback_transcript),
                }
            ],
        )

        # Parse JSON responses
        try:
            strengths_text = strengths_response.content[0].text
            areas_text = areas_response.content[0].text

            # Extract JSON from strengths response
            try:
                strengths_data = json.loads(strengths_text)
            except json.JSONDecodeError:
                start = strengths_text.find("{")
                end = strengths_text.rfind("}") + 1
                if start >= 0 and end > 0:
                    strengths_data = json.loads(strengths_text[start:end])
                else:
                    raise ValueError("No JSON content found in strengths response")

            # Extract JSON from areas response
            try:
                areas_data = json.loads(areas_text)
            except json.JSONDecodeError:
                start = areas_text.find("{")
                end = areas_text.rfind("}") + 1
                if start >= 0 and end > 0:
                    areas_data = json.loads(areas_text[start:end])
                else:
                    raise ValueError("No JSON content found in areas response")

            # Process the strengths data to handle the new structure with is_strong flag
            processed_strengths = {}
            for person, data in strengths_data.get("strengths", {}).items():
                processed_feedback = []
                for feedback_item in data.get("feedback", []):
                    if isinstance(feedback_item, dict):
                        # New format with is_strong flag
                        processed_feedback.append({
                            "text": feedback_item.get("text", ""),
                            "is_strong": feedback_item.get("is_strong", False)
                        })
                    else:
                        # Handle legacy format (plain string)
                        processed_feedback.append({
                            "text": feedback_item,
                            "is_strong": False
                        })
                
                processed_strengths[person] = {
                    "role": data.get("role", ""),
                    "feedback": processed_feedback
                }
            
            # Process the areas_to_target data to handle the new structure with is_strong flag
            processed_areas = {}
            for person, data in areas_data.get("areas_to_target", {}).items():
                processed_feedback = []
                for feedback_item in data.get("feedback", []):
                    if isinstance(feedback_item, dict):
                        # New format with is_strong flag
                        processed_feedback.append({
                            "text": feedback_item.get("text", ""),
                            "is_strong": feedback_item.get("is_strong", False)
                        })
                    else:
                        # Handle legacy format (plain string)
                        processed_feedback.append({
                            "text": feedback_item,
                            "is_strong": False
                        })
                
                processed_areas[person] = {
                    "role": data.get("role", ""),
                    "feedback": processed_feedback
                }

            # Combine results
            result = {
                "strengths": processed_strengths,
                "areas_to_target": processed_areas,
            }

            # Save to cache
            save_cached_data("feedback", file_id, result)

            return result

        except Exception as e:
            print(f"Error parsing responses: {str(e)}")
            print(f"Strengths response: {strengths_text}")
            print(f"Areas response: {areas_text}")
            raise HTTPException(
                status_code=500, detail="Failed to parse AI responses into valid JSON"
            )

    except Exception as e:
        print(f"Error in get_feedback: {str(e)}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/get_development_areas/{file_id}")
async def get_development_areas(
    file_id: str,
    numCompetencies: int,
    use_cache: bool = Query(
        True, description="Whether to use cached results if available"
    ),
    current_user: Dict[str, Any] = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    try:
        # Check cache first with parameters
        if use_cache:
            cached_data = get_cached_data(
                "development_areas", file_id, {"num_competencies": numCompetencies}
            )
            if cached_data:
                print(
                    f"Using cached development areas for file ID {file_id} with {numCompetencies} competencies"
                )
                return cached_data

        # Get previously generated areas to target from the database
        user_id = current_user.user_id if current_user else None
        feedback_data = None
        
        if user_id:
            try:
                feedback_data = get_cached_feedback(user_id, file_id, db)
            except Exception as e:
                print(f"Error getting cached feedback: {str(e)}")
        
        if not feedback_data:
            # If not in database, we need to generate it first
            raise HTTPException(
                status_code=400,
                detail="Feedback data not found. Please generate feedback data first."
            )
            
        print(f"Step 1: Generating competency headings for development areas")
        
        # Step 1: Generate competency headings using the new simplified prompt
        development_prompt = load_prompt("development_headings.txt")
        prompt = development_prompt.format(num_competencies=numCompetencies)
        
        # Generate headings using Claude
        client = anthropic.Anthropic(api_key=api_key)
        response = client.messages.create(
            model="claude-3-7-sonnet-latest",
            max_tokens=1000,
            temperature=0,
            messages=[{"role": "user", "content": prompt}],
        )
        
        # Parse JSON response
        response_text = response.content[0].text
        try:
            headings_result = parse_claude_response(response_text)
        except Exception as e:
            print(f"Error parsing headings response: {str(e)}")
            print(f"Raw response: {response_text}")
            raise HTTPException(
                status_code=500,
                detail="Failed to parse AI response into valid JSON",
            )
        
        # Extract headings as a simple list
        headings = headings_result.get("headings", [])
        print(f"Generated {len(headings)} headings: {headings}")
        
        # Step 2: Use sort_areas_evidence to organize evidence under these headings
        print(f"Step 2: Sorting evidence under generated headings")
        sort_request = SortEvidenceRequest(file_id=file_id, headings=headings)
        sorted_result = await sort_areas_evidence(sort_request, current_user, db)
        
        # Step 3: Transform the result to match the expected frontend format
        print(f"Step 3: Transforming result to match expected frontend format")
        transformed_result = transform_area_evidence(sorted_result)
        
        # Save to cache with parameters
        save_cached_data(
            "development_areas", file_id, transformed_result, {"num_competencies": numCompetencies}
        )
        
        return transformed_result
    except Exception as e:
        print(f"Error in get_development_areas: {str(e)}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/sort-areas-evidence")
async def sort_areas_evidence(
    request: SortEvidenceRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    try:
        # Get previously generated areas to target from the database
        user_id = current_user.user_id
        feedback_data = None
        
        if user_id:
            try:
                feedback_data = get_cached_feedback(user_id, request.file_id, db)
            except Exception as e:
                print(f"Error getting cached feedback: {str(e)}")
        
        if not feedback_data:
            raise HTTPException(
                status_code=400,
                detail="Feedback data not found. Please generate feedback data first."
            )
            
        # Use the previously generated areas to target from the database
        areas_data = feedback_data.get("areas_to_target", {})
        
        # Process the areas data to ensure it has the is_strong flag
        processed_areas = {}
        for person, data in areas_data.items():
            processed_feedback = []
            for feedback_item in data.get("feedback", []):
                if isinstance(feedback_item, dict):
                    # New format with is_strong flag
                    processed_feedback.append({
                        "text": feedback_item.get("text", ""),
                        "is_strong": feedback_item.get("is_strong", False)
                    })
                else:
                    # Handle legacy format (plain string)
                    processed_feedback.append({
                        "text": feedback_item,
                        "is_strong": False
                    })
            
            processed_areas[person] = {
                "role": data.get("role", ""),
                "feedback": processed_feedback
            }
        
        # Process in batches
        all_results = []
        stakeholders = list(processed_areas.keys())
        batch_size = 2  # Process 2 stakeholders at a time (reduced from 3 for better reliability)
        
        # Count total feedback items for verification
        total_feedback_items = sum(len(data.get("feedback", [])) for data in processed_areas.values())
        print(f"Total feedback items to process: {total_feedback_items}")
        
        # Initialize Claude client
        client = anthropic.Anthropic(api_key=api_key)
        
        # Load prompt template
        sort_prompt = load_prompt("sort_evidence_area_to_target.txt")
        
        # Process in batches
        for i in range(0, len(stakeholders), batch_size):
            batch_stakeholders = stakeholders[i:i+batch_size]
            batch_areas = {k: processed_areas[k] for k in batch_stakeholders if k in processed_areas}
            
            # Count items in this batch
            batch_items = sum(len(data.get("feedback", [])) for data in batch_areas.values())
            print(f"Processing batch {i//batch_size + 1} with {len(batch_stakeholders)} stakeholders and {batch_items} feedback items")
            
            # Convert batch data to JSON
            batch_json = json.dumps(batch_areas, indent=2)
            
            # Format prompt for this batch
            prompt = sort_prompt.format(
                areas=batch_json, 
                headings="\n".join(request.headings)
            )
            
            # Get sorted evidence from Claude for this batch
            response = client.messages.create(
                model="claude-3-7-sonnet-latest",
                max_tokens=2000,
                temperature=0,
                messages=[{"role": "user", "content": prompt}],
            )
            
            # Parse JSON response using our enhanced parser
            response_text = response.content[0].text
            
            try:
                # Use our enhanced parser that can handle common JSON issues
                batch_result = parse_claude_response(response_text)
            except Exception as e:
                print(f"Error parsing response: {str(e)}")
                print(f"Raw response: {response_text}")
                raise HTTPException(
                    status_code=500,
                    detail="Failed to parse AI response into valid JSON",
                )
            
            # Count evidence items in the result
            evidence_count = sum(len(item.get("evidence", [])) for item in batch_result)
            print(f"Batch {i//batch_size + 1} returned {evidence_count} evidence items")
            
            # Verify all items were included
            if evidence_count < batch_items:
                print(f"WARNING: Batch {i//batch_size + 1} is missing items. Expected {batch_items}, got {evidence_count}")
            
            all_results.append(batch_result)
        
        # Merge results from all batches
        result = merge_sorted_evidence(all_results)
        
        # Verify total evidence count
        total_evidence = sum(len(item.get("evidence", [])) for item in result)
        print(f"Total evidence items after merging: {total_evidence}")
        
        if total_evidence < total_feedback_items:
            print(f"WARNING: Final result is missing items. Expected {total_feedback_items}, got {total_evidence}")
        
        return result
    except Exception as e:
        print(f"Error in sort_areas_evidence: {str(e)}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))


@app.options("/api/excel")
async def excel_options():
    headers = {
        "Access-Control-Allow-Origin": "http://localhost:3000",
        "Access-Control-Allow-Methods": "GET, OPTIONS",
        "Access-Control-Allow-Headers": "*",
        "Access-Control-Allow-Credentials": "true",
    }
    return Response(content="", headers=headers)


@app.get("/api/excel")
async def get_excel_file():
    file_path = "../Developmental Suggestions & Resources.xlsx"
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Excel file not found")

    with open(file_path, "rb") as f:
        content = f.read()

    headers = {
        "Content-Type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "Content-Disposition": "attachment; filename=Developmental Suggestions & Resources.xlsx",
        "Access-Control-Allow-Origin": "http://localhost:3000",
        "Access-Control-Allow-Methods": "GET, OPTIONS",
        "Access-Control-Allow-Headers": "*",
        "Access-Control-Allow-Credentials": "true",
    }

    return Response(content=content, headers=headers)


@app.post("/api/generate_strength_content")
async def generate_strength_content(
    request: GenerateContentRequest, 
):
    try:
        print(
            f"[Strength Content] Received request - heading: '{request.heading}', file_id: {request.file_id}, has_existing_content: {request.existing_content is not None}"
        )

        # Load transcript using file ID
        feedback_path = f"../data/processed_assessments/filtered_{request.file_id}.txt"
        if not os.path.exists(feedback_path):
            raise HTTPException(
                status_code=404,
                detail=f"Feedback transcript not found at {feedback_path}",
            )

        with open(feedback_path, "r") as f:
            feedback_transcript = f.read()

        # Generate content using Claude
        client = anthropic.Anthropic(api_key=api_key)
        name = get_name_from_report(request.file_id)
        prompt = format_strength_content_prompt(
            name, request.heading, feedback_transcript, request.existing_content
        )

        response = client.messages.create(
            model="claude-3-7-sonnet-latest",
            max_tokens=1000,
            temperature=0,
            messages=[{"role": "user", "content": prompt}],
        )

        return {"content": response.content[0].text}
    except Exception as e:
        print(f"Error in generate_strength_content: {str(e)}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))


app.include_router(file_routers)
app.include_router(feedback_routers)
app.include_router(advice_routers)
app.include_router(snapshot_routers)

if __name__ == "__main__":
    validate_required_env()
    # conn = get_db_connection()

    # Start the API server
    uvicorn.run(
        "main:app", host="0.0.0.0", port=8000, reload=env_variables.RELOAD_MODE
    )
